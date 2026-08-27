from __future__ import annotations

from collections.abc import Mapping
from io import BytesIO
from typing import Any

from docx import Document


def render_docx(snapshot: Mapping[str, Any]) -> bytes:
    """Render exclusively from the persisted itinerary snapshot."""
    document = Document()
    document.add_heading(_text(snapshot.get("title"), "Itinerary"), level=0)
    start_date, end_date = _text(snapshot.get("start_date")), _text(snapshot.get("end_date"))
    if start_date or end_date:
        document.add_paragraph(" - ".join(value for value in (start_date, end_date) if value))
    for number, day in enumerate(_list(snapshot.get("days")), start=1):
        if not isinstance(day, Mapping):
            continue
        document.add_heading(f"Day {number}: {_text(day.get('day_date'))}", level=1)
        for event in _list(day.get("events")):
            if not isinstance(event, Mapping):
                continue
            poi = event.get("poi_snapshot")
            poi_name = _text(poi.get("name")) if isinstance(poi, Mapping) else ""
            title = _text(event.get("notes")) or poi_name or "Activity"
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(title)
            details = " | ".join(value for value in (_text(event.get("starts_at")), _text(event.get("ends_at"))) if value)
            if details:
                paragraph.add_run(f" ({details})")
            if poi_name and poi_name != title:
                document.add_paragraph(poi_name)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _text(value: object, default: str = "") -> str:
    return value.strip() if isinstance(value, str) else default


def _list(value: object) -> list[object]:
    return value if isinstance(value, list | tuple) else []
